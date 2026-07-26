from docx import Document
from docx.shared import Pt, Cm
import pyautogui
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import ipaddress
import os
from datetime import datetime, timezone

# ==================== НАСТРОЙКИ ====================
# Введите VirusTotal API Key и AbuseIPDB API Key
VT_API_KEY = os.getenv("VT_API_KEY")
ABUSEIPDB_API_KEY = os.getenv("ABUSEIPDB_API_KEY")
if not VT_API_KEY:
    raise ValueError("VT_API_KEY не задана в переменных окружения")
if not ABUSEIPDB_API_KEY:
    raise ValueError("ABUSEIPDB_API_KEY не задана в переменных окружения")

MAX_AGE_IN_DAYS = "90"

# Создание окошка с запросом пользовательского ввода
user_input = pyautogui.prompt(text='Введите ip-адреса + страны', title='Отчёт о вредоносных IP-адресах' , default='')

# Модификация пользовательского ввода для правильного форматирования данных
user_input = user_input.split('\n')
modified_user_input = []
for elem in user_input:
    elem = elem.replace(' ', '@', 1)
    elem = elem + '@'
    modified_user_input.append(elem)
modified_user_input = ''.join(modified_user_input).split('@')
modified_user_input.pop()
user_input = modified_user_input

# Создание документа
doc = Document()

# Задание стилей документа
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.15

# Получение текущей даты
now_date = str(datetime.now().strftime("%d-%m-%Y %H:%M").replace(':','-'))

# Вывод в документ текста с датой
doc.add_paragraph(f'Уведомление об автоматизированных сканированиях {now_date.split()[0].replace('-', '.')}')

# Создание таблицы в документе и задание стилей
table = doc.add_table(rows=len(user_input)//2+1, cols=2)
table.style = 'Table Grid'
table.autofit = False
for row in table.rows:
    row.cells[0].width = Cm(2.94)
    row.cells[1].width = Cm(13.74)
table.rows[0].height = Cm(0.62)

# Задание заголовков таблицы
cell = table.cell(0, 0)
cell.text = 'IP'
cell.paragraphs[0].runs[0].bold = True
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

cell = table.cell(0, 1)
cell.text = 'Описание'
cell.paragraphs[0].runs[0].bold = True
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Вывод айпишников и их стран
i = 0
for row in range(1 ,len(user_input)//2+1):
    cell = table.cell(row, 0)
    ip_parts = user_input[i].rsplit('.', 1)
    cell.text = ip_parts[0] + '*' + ip_parts[1] + f'\n({user_input[i+1]})'
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i += 2

# Объединение всех строк во втором столбце
if len(user_input)//2 > 0:
    start_cell = table.cell(1,1)
    end_cell = table.cell(len(user_input)//2, 1)
    start_cell.merge(end_cell)

# Проверка всех айпи адресов и выдача самого вредоносного ##############################################################################################

# ==================== СЛОВАРИ ====================

# Словарь для преобразования кодов стран в названия на русском ПИСЬКА с===3
country_codes_to_russian = {
    "AF": "Афганистан", "AX": "Аландские острова", "AL": "Албания", "DZ": "Алжир",
    "AS": "Американское Самоа", "AD": "Андорра", "AO": "Ангола", "AI": "Ангилья",
    "AQ": "Антарктида", "AG": "Антигуа и Барбуда", "AR": "Аргентина", "AM": "Армения",
    "AW": "Аруба", "AU": "Австралия", "AT": "Австрия", "AZ": "Азербайджан",
    "BS": "Багамы", "BH": "Бахрейн", "BD": "Бангладеш", "BB": "Барбадос",
    "BY": "Беларусь", "BE": "Бельгия", "BZ": "Белиз", "BJ": "Бенин",
    "BM": "Бермуды", "BT": "Бутан", "BO": "Боливия", "BQ": "Бонайре, Синт-Эстатиус и Саба",
    "BA": "Босния и Герцеговина", "BW": "Ботсвана", "BV": "Остров Буве", "BR": "Бразилия",
    "IO": "Британская территория в Индийском океане", "BN": "Бруней-Даруссалам", "BG": "Болгария",
    "BF": "Буркина-Фасо", "BI": "Бурунди", "CV": "Кабо-Верде", "KH": "Камбоджа",
    "CM": "Камерун", "CA": "Канада", "KY": "Острова Кайман", "CF": "ЦАР",
    "TD": "Чад", "CL": "Чили", "CN": "Китай", "CX": "Остров Рождества",
    "CC": "Кокосовые острова", "CO": "Колумбия", "KM": "Коморы", "CG": "Конго",
    "CD": "ДР Конго", "CK": "Острова Кука", "CR": "Коста-Рика", "CI": "Кот-д'Ивуар",
    "HR": "Хорватия", "CU": "Куба", "CW": "Кюрасао", "CY": "Кипр",
    "CZ": "Чехия", "DK": "Дания", "DJ": "Джибути", "DM": "Доминика",
    "DO": "Доминиканская Республика", "EC": "Эквадор", "EG": "Египет", "SV": "Сальвадор",
    "GQ": "Экваториальная Гвинея", "ER": "Эритрея", "EE": "Эстония", "SZ": "Эсватини",
    "ET": "Эфиопия", "FK": "Фолклендские острова", "FO": "Фарерские острова", "FJ": "Фиджи",
    "FI": "Финляндия", "FR": "Франция", "GF": "Французская Гвиана", "PF": "Французская Полинезия",
    "TF": "Французские Южные территории", "GA": "Габон", "GM": "Гамбия", "GE": "Грузия",
    "DE": "Германия", "GH": "Гана", "GI": "Гибралтар", "GR": "Греция",
    "GL": "Гренландия", "GD": "Гренада", "GP": "Гваделупа", "GU": "Гуам",
    "GT": "Гватемала", "GG": "Гернси", "GN": "Гвинея", "GW": "Гвинея-Бисау",
    "GY": "Гайана", "HT": "Гаити", "HM": "Херд и Макдональд", "VA": "Ватикан",
    "HN": "Гондурас", "HK": "Гонконг", "HU": "Венгрия", "IS": "Исландия",
    "IN": "Индия", "ID": "Индонезия", "IR": "Иран", "IQ": "Ирак",
    "IE": "Ирландия", "IM": "Остров Мэн", "IL": "Израиль", "IT": "Италия",
    "JM": "Ямайка", "JP": "Япония", "JE": "Джерси", "JO": "Иордания",
    "KZ": "Казахстан", "KE": "Кения", "KI": "Кирибати", "KP": "КНДР",
    "KR": "Республика Корея", "KW": "Кувейт", "KG": "Киргизия", "LA": "Лаос",
    "LV": "Латвия", "LB": "Ливан", "LS": "Лесото", "LR": "Либерия",
    "LY": "Ливия", "LI": "Лихтенштейн", "LT": "Литва", "LU": "Люксембург",
    "MO": "Макао", "MG": "Мадагаскар", "MW": "Малави", "MY": "Малайзия",
    "MV": "Мальдивы", "ML": "Мали", "MT": "Мальта", "MH": "Маршалловы Острова",
    "MQ": "Мартиника", "MR": "Мавритания", "MU": "Маврикий", "YT": "Майотта",
    "MX": "Мексика", "FM": "Микронезия", "MD": "Молдавия", "MC": "Монако",
    "MN": "Монголия", "ME": "Черногория", "MS": "Монтсеррат", "MA": "Марокко",
    "MZ": "Мозамбик", "MM": "Мьянма", "NA": "Намибия", "NR": "Науру",
    "NP": "Непал", "NL": "Нидерланды", "NC": "Новая Каледония", "NZ": "Новая Зеландия",
    "NI": "Никарагуа", "NE": "Нигер", "NG": "Нигерия", "NU": "Ниуэ",
    "NF": "Остров Норфолк", "MK": "Северная Македония", "MP": "Северные Марианские Острова",
    "NO": "Норвегия", "OM": "Оман", "PK": "Пакистан", "PW": "Палау",
    "PS": "Палестина", "PA": "Панама", "PG": "Папуа — Новая Гвинея", "PY": "Парагвай",
    "PE": "Перу", "PH": "Филиппины", "PN": "Острова Питкэрн", "PL": "Польша",
    "PT": "Португалия", "PR": "Пуэрто-Рико", "QA": "Катар", "RE": "Реюньон",
    "RO": "Румыния", "RU": "Россия", "RW": "Руанда", "BL": "Сен-Бартелеми",
    "SH": "Острова Святой Елены, Вознесения и Тристан-да-Кунья", "KN": "Сент-Китс и Невис",
    "LC": "Сент-Люсия", "MF": "Сен-Мартен", "PM": "Сен-Пьер и Микелон", "VC": "Сент-Винсент и Гренадины",
    "WS": "Самоа", "SM": "Сан-Марино", "ST": "Сан-Томе и Принсипи", "SA": "Саудовская Аравия",
    "SN": "Сенегал", "RS": "Сербия", "SC": "Сейшелы", "SL": "Сьерра-Леоне",
    "SG": "Сингапур", "SX": "Синт-Мартен", "SK": "Словакия", "SI": "Словения",
    "SB": "Соломоновы Острова", "SO": "Сомали", "ZA": "ЮАР", "GS": "Южная Георгия и Южные Сандвичевы острова",
    "SS": "Южный Судан", "ES": "Испания", "LK": "Шри-Ланка", "SD": "Судан",
    "SR": "Суринам", "SJ": "Шпицберген и Ян-Майен", "SE": "Швеция", "CH": "Швейцария",
    "SY": "Сирия", "TW": "Тайвань", "TJ": "Таджикистан", "TZ": "Танзания",
    "TH": "Таиланд", "TL": "Восточный Тимор", "TG": "Того", "TK": "Токелау",
    "TO": "Тонга", "TT": "Тринидад и Тобаго", "TN": "Тунис", "TR": "Турция",
    "TM": "Туркменистан", "TC": "Тёркс и Кайкос", "TV": "Тувалу", "UG": "Уганда",
    "UA": "Украина", "AE": "ОАЭ", "GB": "Великобритания", "US": "США",
    "UM": "Внешние малые острова США", "UY": "Уругвай", "UZ": "Узбекистан", "VU": "Вануату",
    "VE": "Венесуэла", "VN": "Вьетнам", "VG": "Виргинские Острова (Великобритания)",
    "VI": "Виргинские Острова (США)", "WF": "Уоллис и Футуна", "EH": "Западная Сахара", "YE": "Йемен",
    "ZM": "Замбия", "ZW": "Зимбабве", "XK": "Косово",
}

# Категории нарушений AbuseIPDB (https://www.abuseipdb.com/categories)
# Названия оставлены в исходном виде, как они фигурируют на самой платформе AbuseIPDB
abuseipdb_categories_en = {
    1: "DNS Compromise",
    2: "DNS Poisoning",
    3: "Fraud Orders",
    4: "DDoS Attack",
    5: "FTP Brute-Force",
    6: "Ping of Death",
    7: "Phishing",
    8: "Fraud VOIP",
    9: "Open Proxy",
    10: "Web Spam",
    11: "Email Spam",
    12: "Blog Spam",
    13: "VPN IP",
    14: "Port Scan",
    15: "Hacking",
    16: "SQL Injection",
    17: "Spoofing",
    18: "Brute-Force",
    19: "Bad Web Bot",
    20: "Exploited Host",
    21: "Web App Attack",
    22: "SSH",
    23: "IoT Targeted",
}


def get_country_name_russian(country_code):
    """Преобразует код страны в название на русском языке"""
    return country_codes_to_russian.get(country_code, f"{country_code} (неизвестная страна)")


def get_category_name(category_id):
    """Преобразует числовой ID категории AbuseIPDB в её короткое название"""
    return abuseipdb_categories_en.get(category_id, f"Unknown ({category_id})")


def ru_plural(number: int, one: str, few: str, many: str) -> str:
    """Подбирает правильную форму слова для числительного (1 час / 2 часа / 5 часов)"""
    n = abs(number) % 100
    n1 = n % 10
    if 11 <= n <= 14:
        return many
    if n1 == 1:
        return one
    if 2 <= n1 <= 4:
        return few
    return many


def parse_iso_datetime(value: str | None) -> datetime | None:
    """Парсит ISO 8601 дату из ответа AbuseIPDB в datetime с таймзоной UTC"""
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt


def humanize_elapsed(dt: datetime) -> str:
    """Форматирует промежуток времени от dt до сейчас, например '4 часа'"""
    seconds = max(int((datetime.now(timezone.utc) - dt).total_seconds()), 0)

    if seconds < 60:
        return "несколько минут"

    minutes = seconds // 60
    if minutes < 60:
        return f"{minutes} {ru_plural(minutes, 'минуту', 'минуты', 'минут')}"

    hours = minutes // 60
    if hours < 24:
        return f"{hours} {ru_plural(hours, 'час', 'часа', 'часов')}"

    days = hours // 24
    return f"{days} {ru_plural(days, 'день', 'дня', 'дней')}"


def is_valid_ip(value: str) -> bool:
    """Проверяет, что строка — корректный IPv4 или IPv6 адрес"""
    try:
        ipaddress.ip_address(value)
        return True
    except ValueError:
        return False


# ==================== ЗАПРОСЫ К API ====================

def fetch_virustotal(ip_address: str):
    url = f"https://www.virustotal.com/api/v3/ip_addresses/{ip_address}"
    headers = {
        "accept": "application/json",
        "x-apikey": VT_API_KEY
    }

    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        return response.status_code, response.json()

    return response.status_code, response.text


def fetch_abuseipdb(ip_address: str):
    url = "https://api.abuseipdb.com/api/v2/check"

    params = {
        "ipAddress": ip_address,
        "maxAgeInDays": MAX_AGE_IN_DAYS,
        "verbose": "true",
    }

    headers = {
        "Accept": "application/json",
        "Key": ABUSEIPDB_API_KEY
    }

    response = requests.get(
        url,
        headers=headers,
        params=params
    )

    if response.status_code == 200:
        return response.status_code, response.json()

    return response.status_code, response.text


# ==================== ФОРМИРОВАНИЕ ОТЧЁТА ====================

def build_report(ip_address: str) -> str:

    vt_status, vt_data = fetch_virustotal(ip_address)
    abuse_status, abuse_data = fetch_abuseipdb(ip_address)

    report_lines = [f"- IP-адрес: {ip_address}"]
    abuse_categories_line = None 

    # ---------- AbuseIPDB: страна и провайдер ----------
    if abuse_status == 200 and isinstance(abuse_data, dict):
        data = abuse_data.get("data", {})

        country_code = data.get("countryCode")
        isp = data.get("isp") or "не определён"

        if country_code:
            report_lines.append(f"Страна: {get_country_name_russian(country_code)}")
        else:
            report_lines.append("Страна: не определена")

        report_lines.append(f"Провайдер: {isp}")

        # Собираем уникальные категории из всех отчётов пользователей,
        # а также даты этих отчётов, чтобы найти самый старый (первый) отчёт
        category_ids = set()
        report_dates = []
        for report in data.get("reports", []):
            category_ids.update(report.get("categories", []))
            parsed_date = parse_iso_datetime(report.get("reportedAt"))
            if parsed_date:
                report_dates.append(parsed_date)

        if category_ids:
            categories_text = ", ".join(
                get_category_name(cid) for cid in sorted(category_ids)
            )
            oldest_report_at = min(report_dates) if report_dates else None
            time_ago = humanize_elapsed(oldest_report_at) if oldest_report_at else f"{MAX_AGE_IN_DAYS} дней"
            abuse_categories_line = (
                f"По отчётам пользователей на AbuseIPDB за последние {time_ago} "
                f"представлен в категориях: {categories_text}"
            )
    else:
        report_lines.append(f"Не удалось получить данные AbuseIPDB (код {abuse_status}).")

    # ---------- VirusTotal ----------
    if vt_status == 200 and isinstance(vt_data, dict):
        suspicious, phishing, malicious, malware = [], [], [], []

        attributes = vt_data.get("data", {}).get("attributes", {})
        stats = attributes.get("last_analysis_stats", {})
        total_vendors = sum(stats.values()) if stats else 0

        results = attributes.get("last_analysis_results", {})
        for engine_name, engine_data in results.items():
            result = (engine_data.get("result") or "").lower()
            if "suspicious" in result:
                suspicious.append(engine_name)
            elif "phishing" in result:
                phishing.append(engine_name)
            elif "malicious" in result:
                malicious.append(engine_name)
            elif "malware" in result:
                malware.append(engine_name)

        total_detections = len(suspicious) + len(phishing) + len(malicious) + len(malware)

        if total_detections == 0:
            report_lines.append(
                "При проверке на TI платформах не установлено фактов вредоносной активности IP-адреса."
            )
        else:
            report_lines.append(
                "При проверке на TI платформах установлено, что IP-адрес был зафиксирован во вредоносной "
                f"активности. \nПо данным VirusTotal {total_detections}/{total_vendors} TI вендоров отметили "
                "IP-адрес в своих отчётах:"
            )
            for label, vendors in (
                ("Suspicious", suspicious),
                ("Phishing", phishing),
                ("Malicious", malicious),
                ("Malware", malware),
            ):
                if vendors:
                    report_lines.append(f'С классификацией "{label}":')
                    for vendor in vendors:
                        report_lines.append(f"    - {vendor}")
    else:
        report_lines.append(f"Ошибка при запросе к VirusTotal (код {vt_status}).")

    # ---------- AbuseIPDB ----------
    if abuse_categories_line:
        report_lines.append(abuse_categories_line)

    return "\n".join(report_lines), total_detections

#####################################################################################################################################################

# Пробиваем все айпишники и сохраняем репорты на них
ip_on_check = user_input[::2]
map_detections_report = {}
for ip_address in ip_on_check:
    report, detections = build_report(ip_address)
    if detections not in map_detections_report:
        map_detections_report[detections] = report

# Находим самый вредоносный айпишник
map_detections_report = dict(sorted(map_detections_report.items(), reverse=True))
max_detections_report = list(map_detections_report.values())[0]

# Преобразуем отчёт о наивреднейшем айпи для отчёта в docx
max_detections_report = max_detections_report[max_detections_report.find("При проверке на TI"):]
max_detections_report = max_detections_report.replace('IP-адрес был зафиксирован', 'IP-адреса были зафиксированы', 1)
max_detections_report = max_detections_report.replace('IP-адрес в', 'IP-адреса в', 1)
max_detections_report = max_detections_report.replace('представлен в', 'представлены в', 1)

# Выводим вреднейший отчёт в таблицу 
cell = table.cell(1, 1)
cell.text = max_detections_report

# Сохранение документа с выбранным названием
doc.save(f'Отчёт о сканах {now_date}.docx')
print('Отчёт готов. Завершение работы.')